#!/usr/bin/env python3
"""
Search through downloaded YouTube comments for keywords and patterns
"""

import json
import argparse
import re
import sys
from pathlib import Path
from collections import defaultdict
from datetime import datetime

# By default, ANSI colors are enabled
USE_COLOR = True





APP_NAME = "YouTube Comment Search"

# Force UTF-8 output on Windows
if sys.platform.startswith("win"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except AttributeError:
        pass  # Python < 3.7 fallback

# ANSI color codes for terminal highlighting
RED = '\033[91m'
GREEN = '\033[92m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
RESET = '\033[0m'

class CommentSearcher:
    def __init__(self, comments_dir):
        self.comments_dir = Path(comments_dir)
        self.results = []
        
    def load_comments_from_file(self, filepath):
        """Load yt-dlp comments from a .info.json file"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)

            # yt-dlp stores comments under data["comments"]
            comments = data.get("comments", [])
            if not isinstance(comments, list):
                return [], filepath.name

            return comments, filepath.name

        except Exception as e:
            print(f"✗ Error loading {filepath}: {e}")
            return [], None


    def extract_video_id_from_filename(self, filename: str) -> str:
        """Extract video ID from <videoid>.info.json"""
        if filename.endswith(".info.json"):
            return filename[:-len(".info.json")]
        return "unknown"

    
    def search_simple(self, text, keywords, case_sensitive=False):
        """Simple keyword search (any keyword matches)"""
        if not case_sensitive:
            text = text.lower()
            keywords = [k.lower() for k in keywords]
        
        return any(keyword in text for keyword in keywords)
    
    def search_all_keywords(self, text, keywords, case_sensitive=False):
        """Search requiring ALL keywords to be present"""
        if not case_sensitive:
            text = text.lower()
            keywords = [k.lower() for k in keywords]
        
        return all(keyword in text for keyword in keywords)
    
    def search_phrase(self, text, phrase, case_sensitive=False):
        """Search for exact phrase"""
        if not case_sensitive:
            text = text.lower()
            phrase = phrase.lower()
        
        return phrase in text
    
    def search_regex(self, text, pattern):
        """Search using regular expression"""
        try:
            return bool(re.search(pattern, text, re.IGNORECASE))
        except re.error as e:
            print(f"✗ Invalid regex pattern: {e}")
            return False
    
    def highlight_matches(self, text, keywords, search_mode, case_sensitive=False, use_color=True):
        """Highlight matching keywords in text"""
        if not use_color:
            return text
        
        highlighted_text = text
        
        if search_mode == 'regex':
            # For regex, highlight actual matched portions
            try:
                pattern = re.compile(keywords[0], re.IGNORECASE if not case_sensitive else 0)
                matches = list(pattern.finditer(text))
                # Process matches in reverse to maintain string positions
                for match in reversed(matches):
                    start, end = match.span()
                    highlighted_text = (highlighted_text[:start] + 
                                      RED + highlighted_text[start:end] + RESET + 
                                      highlighted_text[end:])
            except re.error:
                pass
        elif search_mode == 'phrase':
            # Highlight the entire phrase
            phrase = ' '.join(keywords)
            if not case_sensitive:
                # Case-insensitive phrase highlighting
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                highlighted_text = pattern.sub(lambda m: RED + m.group(0) + RESET, highlighted_text)
            else:
                highlighted_text = highlighted_text.replace(phrase, RED + phrase + RESET)
        else:
            # Highlight individual keywords (any or all mode)
            for keyword in keywords:
                if not case_sensitive:
                    # Case-insensitive keyword highlighting
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    highlighted_text = pattern.sub(lambda m: RED + m.group(0) + RESET, highlighted_text)
                else:
                    highlighted_text = highlighted_text.replace(keyword, RED + keyword + RESET)
        
        return highlighted_text
    
    def get_plain_matches(self, text, keywords, search_mode, case_sensitive=False):
        """Get list of actual matched text portions (for DOCX highlighting)"""
        matches = []
        
        if search_mode == 'regex':
            try:
                pattern = re.compile(keywords[0], re.IGNORECASE if not case_sensitive else 0)
                for match in pattern.finditer(text):
                    matches.append(match.group(0))
            except re.error:
                pass
        elif search_mode == 'phrase':
            phrase = ' '.join(keywords)
            if not case_sensitive:
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                matches = [m.group(0) for m in pattern.finditer(text)]
            else:
                # Find all occurrences of exact phrase
                start = 0
                while True:
                    pos = text.find(phrase, start)
                    if pos == -1:
                        break
                    matches.append(text[pos:pos+len(phrase)])
                    start = pos + 1
        else:
            for keyword in keywords:
                if not case_sensitive:
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    matches.extend([m.group(0) for m in pattern.finditer(text)])
                else:
                    start = 0
                    while True:
                        pos = text.find(keyword, start)
                        if pos == -1:
                            break
                        matches.append(text[pos:pos+len(keyword)])
                        start = pos + 1
        
        return list(set(matches))  # Remove duplicates
    
    def calculate_relevance_score(self, comment, keywords, case_sensitive=False):
        """Calculate relevance score based on keyword frequency and position"""
        text = comment['text']
        if not case_sensitive:
            text_lower = text.lower()
            keywords_lower = [k.lower() for k in keywords]
        else:
            text_lower = text
            keywords_lower = keywords
        
        score = 0
        
        # Count keyword occurrences
        for keyword in keywords_lower:
            count = text_lower.count(keyword)
            score += count * 10
        
        # Bonus for keyword in first 50 characters
        if any(kw in text_lower[:50] for kw in keywords_lower):
            score += 5
        
        # Bonus for higher like count
        score += min(comment.get('like_count', 0), 50)
        
        # Bonus for uploader/verified authors
        if comment.get('author_is_uploader'):
            score += 20
        if comment.get('author_is_verified'):
            score += 10
        
        return score
    
    def get_all_users(self):
        print(f"\n{'='*60}")
        print(f"Analyzing users in: {self.comments_dir}")
        print(f"{'='*60}\n")

        if not self.comments_dir.exists():
            print(f"✗ Error: Directory not found: {self.comments_dir}")
            return {}

        json_files = list(self.comments_dir.glob("*.info.json"))

        if not json_files:
            print(f"✗ No *.info.json files found in {self.comments_dir}")
            return {}

        print(f"Scanning {len(json_files)} info JSON file(s)...\n")

        user_stats = defaultdict(lambda: {
            "count": 0,
            "total_likes": 0,
            "comments": [],
            "author_id": None,
            "is_verified": False,
            "is_uploader": False,
        })

        total_comments = 0

        for json_file in json_files:
            comments, filename = self.load_comments_from_file(json_file)
            video_id = self.extract_video_id_from_filename(filename)

            total_comments += len(comments)

            for comment in comments:
                author = comment.get("author", "Unknown")
                author_id = comment.get("author_id", "N/A")

                stats = user_stats[author]
                stats["count"] += 1
                stats["total_likes"] += comment.get("like_count", 0)
                stats["author_id"] = author_id

                if comment.get("author_is_verified"):
                    stats["is_verified"] = True
                if comment.get("author_is_uploader"):
                    stats["is_uploader"] = True

                stats["comments"].append({
                    "text": comment.get("text", ""),
                    "video_id": video_id,
                    "source_file": filename,
                    "like_count": comment.get("like_count", 0),
                    "timestamp": comment.get("timestamp", 0),
                    "time_text": comment.get("_time_text", "Unknown"),
                    "full_comment": comment,
                })

        print(f"✓ Scanned {total_comments} total comments")
        print(f"✓ Found {len(user_stats)} unique users\n")

        return dict(user_stats)


    
    def show_most_active_users(self, limit=10):
        """Display the most active users"""
        user_stats = self.get_all_users()
        
        if not user_stats:
            return []
        
        # Sort by comment count
        sorted_users = sorted(user_stats.items(), 
                             key=lambda x: x[1]['count'], 
                             reverse=True)
        
        print(f"{'='*60}")
        print(f"Top {min(limit, len(sorted_users))} Most Active Users")
        print(f"{'='*60}\n")
        
        top_users = []
        for i, (author, stats) in enumerate(sorted_users[:limit], 1):
            badges = []
            if stats['is_verified']:
                badges.append(f"{BLUE}✓ Verified{RESET}")
            if stats['is_uploader']:
                badges.append(f"{GREEN}⬆ Uploader{RESET}")
            
            badge_str = f" [{', '.join(badges)}]" if badges else ""
            
            print(f"{YELLOW}{i}. {author}{RESET}{badge_str}")
            print(f"   Author ID: {stats['author_id']}")
            print(f"   Comments: {stats['count']} | Total likes: {stats['total_likes']}")
            print(f"   Avg likes per comment: {stats['total_likes'] / stats['count']:.1f}")
            print()
            
            top_users.append((author, stats))
        
        return top_users
    
    def export_active_users_json(self, top_users, output_file, command_line):
        """Export most active users to JSON file"""
        try:
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            export_data = {
                'command': command_line,
                'generated': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'total_users': len(top_users),
                'users': []
            }
            
            for rank, (author, stats) in enumerate(top_users, 1):
                export_data['users'].append({
                    'rank': rank,
                    'author': author,
                    'author_id': stats['author_id'],
                    'comment_count': stats['count'],
                    'total_likes': stats['total_likes'],
                    'avg_likes': stats['total_likes'] / stats['count'] if stats['count'] > 0 else 0,
                    'is_verified': stats['is_verified'],
                    'is_uploader': stats['is_uploader'],
                })
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            print(f"✓ Most active users exported to JSON: {output_file}")
        except Exception as e:
            print(f"✗ Error exporting JSON: {e}")
    
    def export_active_users_docx(self, top_users, output_file, command_line):
        """Export most active users to DOCX file"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = Document()
            doc.add_heading('Most Active Users Analysis', 0)
            
            if command_line:
                p = doc.add_paragraph()
                p.add_run('Command: ').bold = True
                p.add_run(command_line)
                doc.add_paragraph()
            
            doc.add_paragraph(f'Total Users Analyzed: {len(top_users)}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            for rank, (author, stats) in enumerate(top_users, 1):
                doc.add_heading(f'Rank {rank}: {author}', level=1)
                
                p = doc.add_paragraph()
                p.add_run(f"Author ID: ").bold = True
                p.add_run(f"{stats['author_id']}\n")
                
                p.add_run(f"Comments: ").bold = True
                p.add_run(f"{stats['count']}\n")
                
                p.add_run(f"Total Likes: ").bold = True
                p.add_run(f"{stats['total_likes']}\n")
                
                p.add_run(f"Avg Likes per Comment: ").bold = True
                p.add_run(f"{stats['total_likes'] / stats['count']:.1f}\n")
                
                badges = []
                if stats['is_verified']:
                    badges.append("Verified")
                if stats['is_uploader']:
                    badges.append("Uploader")
                
                if badges:
                    p.add_run(f"Badges: ").bold = True
                    p.add_run(f"{', '.join(badges)}\n")
                
                doc.add_paragraph()
            
            doc.save(output_file)
            print(f"✓ Most active users exported to DOCX: {output_file}")
            
        except ImportError:
            print(f"✗ Error: python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            print(f"✗ Error exporting DOCX: {e}")
    
    def extract_user_comments(self, username, search_keywords=None, search_mode='any', 
                             case_sensitive=False, min_likes=0):
        """Extract all comments from a specific user, optionally filtering by keywords"""
        user_stats = self.get_all_users()
        
        if username not in user_stats:
            print(f"✗ User '{username}' not found in comments")
            print(f"\nDid you mean one of these?")
            
            # Find similar usernames (case-insensitive partial match)
            similar = [u for u in user_stats.keys() if username.lower() in u.lower()]
            for u in similar[:5]:
                print(f"  - {u}")
            
            return []
        
        stats = user_stats[username]
        all_comments = stats['comments']
        
        print(f"\n{'='*60}")
        print(f"User: {GREEN}{username}{RESET}")
        print(f"Author ID: {stats['author_id']}")
        
        badges = []
        if stats['is_verified']:
            badges.append("Verified")
        if stats['is_uploader']:
            badges.append("Uploader")
        if badges:
            print(f"Badges: {', '.join(badges)}")
        
        print(f"Total comments: {stats['count']}")
        print(f"Total likes: {stats['total_likes']}")
        print(f"{'='*60}\n")
        
        # Filter by keywords if provided
        if search_keywords:
            print(f"Filtering by keywords: {search_keywords}")
            print(f"Mode: {search_mode}, Case-sensitive: {case_sensitive}\n")
            
            filtered_comments = []
            for comment_data in all_comments:
                text = comment_data['text']
                
                # Apply like filter
                if comment_data['like_count'] < min_likes:
                    continue
                
                # Apply search based on mode
                match = False
                if search_mode == 'any':
                    match = self.search_simple(text, search_keywords, case_sensitive)
                elif search_mode == 'all':
                    match = self.search_all_keywords(text, search_keywords, case_sensitive)
                elif search_mode == 'phrase':
                    match = self.search_phrase(text, ' '.join(search_keywords), case_sensitive)
                elif search_mode == 'regex':
                    match = self.search_regex(text, search_keywords[0])
                
                if match:
                    filtered_comments.append(comment_data)
            
            all_comments = filtered_comments
            print(f"✓ Found {len(all_comments)} matching comments from this user\n")
        else:
            # Apply like filter even without keywords
            if min_likes > 0:
                all_comments = [c for c in all_comments if c['like_count'] >= min_likes]
                print(f"✓ Filtered to {len(all_comments)} comments with {min_likes}+ likes\n")
        
        # Convert to results format
        results = []
        for comment_data in all_comments:
            result = {
                'comment': comment_data['full_comment'],
                'source_file': comment_data['source_file'],
                'video_id': comment_data['video_id'],
                'relevance_score': 0,  # Not applicable for user extraction
                'matched_text': []
            }
            
            if search_keywords:
                result['relevance_score'] = self.calculate_relevance_score(
                    comment_data['full_comment'], 
                    search_keywords, 
                    case_sensitive
                )
                result['matched_text'] = self.get_plain_matches(
                    comment_data['text'], 
                    search_keywords, 
                    search_mode, 
                    case_sensitive
                )
            
            results.append(result)
        
        # Sort by likes if no keywords, otherwise by relevance
        if search_keywords:
            results.sort(key=lambda x: x['relevance_score'], reverse=True)
        else:
            results.sort(key=lambda x: x['comment'].get('like_count', 0), reverse=True)
        
        return results
    
    def search_comments(
        self,
        keywords,
        search_mode="any",
        case_sensitive=False,
        min_likes=0,
        max_results=None,
        sort_by="relevance",
    ):
        print(f"\n{'='*60}")
        print(f"Searching in: {self.comments_dir}")
        print(f"Keywords: {keywords}")
        print(f"Mode: {search_mode}, Case-sensitive: {case_sensitive}")
        print(f"Min likes: {min_likes}")
        print(f"{'='*60}\n")

        if not self.comments_dir.exists():
            print(f"✗ Error: Directory not found: {self.comments_dir}")
            return []

        json_files = list(self.comments_dir.glob("*.info.json"))

        if not json_files:
            print(f"✗ No *.info.json files found in {self.comments_dir}")
            return []

        print(f"Scanning {len(json_files)} info JSON file(s)...\n")

        results = []
        total_comments = 0

        for json_file in json_files:
            comments, filename = self.load_comments_from_file(json_file)
            video_id = self.extract_video_id_from_filename(filename)

            total_comments += len(comments)

            for comment in comments:
                text = comment.get("text", "")
                if not text:
                    continue

                if comment.get("like_count", 0) < min_likes:
                    continue

                match = False
                if search_mode == "any":
                    match = self.search_simple(text, keywords, case_sensitive)
                elif search_mode == "all":
                    match = self.search_all_keywords(text, keywords, case_sensitive)
                elif search_mode == "phrase":
                    match = self.search_phrase(text, " ".join(keywords), case_sensitive)
                elif search_mode == "regex":
                    match = self.search_regex(text, keywords[0])

                if match:
                    results.append({
                        "comment": comment,
                        "source_file": filename,
                        "video_id": video_id,
                        "relevance_score": self.calculate_relevance_score(
                            comment, keywords, case_sensitive
                        ),
                        "matched_text": self.get_plain_matches(
                            text, keywords, search_mode, case_sensitive
                        ),
                    })

        print(f"✓ Scanned {total_comments} total comments")
        print(f"✓ Found {len(results)} matching comments\n")

        if sort_by == "relevance":
            results.sort(key=lambda x: x["relevance_score"], reverse=True)
        elif sort_by == "likes":
            results.sort(key=lambda x: x["comment"].get("like_count", 0), reverse=True)
        elif sort_by == "date":
            results.sort(key=lambda x: x["comment"].get("timestamp", 0), reverse=True)

        if max_results:
            results = results[:max_results]

        return results


    
    def display_results(self, results, highlight=True, keywords=None, search_mode='any', 
                       case_sensitive=False):
        """Display search results"""
        if not results:
            print("No results found.")
            return
        
        print(f"{'='*60}")
        print(f"Top {len(results)} Results")
        print(f"{'='*60}\n")
        
        for i, result in enumerate(results, 1):
            comment = result['comment']
            
            print(f"[{i}] Video ID: {result['video_id']}")
            print(f"    Author: {comment.get('author', 'Unknown')} ({comment.get('author_id', 'N/A')})")
            print(f"    Likes: {comment.get('like_count', 0)} | "
                  f"Verified: {'Yes' if comment.get('author_is_verified') else 'No'} | "
                  f"Uploader: {'Yes' if comment.get('author_is_uploader') else 'No'}")
            print(f"    Posted: {comment.get('_time_text', 'Unknown')}")
            
            if result['relevance_score'] > 0:
                print(f"    Relevance Score: {result['relevance_score']}")
            
            # Display comment text (never truncated)
            text = comment.get('text', '')
            if highlight and keywords:
                text = self.highlight_matches(text, keywords, search_mode, case_sensitive, use_color=USE_COLOR)

            print(f"    Comment: {text}")
            
            print(f"    URL: https://www.youtube.com/watch?v={result['video_id']}")
            print()
    
    def export_json(self, results, output_file, command_line, username=None):
        """Export results to JSON file"""
        try:
            # Create output directory if it doesn't exist
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            export_data = {
                'command': command_line,
                'generated': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'total_results': len(results),
                'results': []
            }
            
            if username:
                export_data['user_filter'] = username
            
            for result in results:
                export_data['results'].append({
                    'video_id': result['video_id'],
                    'source_file': result['source_file'],
                    'relevance_score': result['relevance_score'],
                    'matched_text': result.get('matched_text', []),
                    'comment': result['comment']
                })
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            print(f"✓ Results exported to JSON: {output_file}")
        except Exception as e:
            print(f"✗ Error exporting JSON results: {e}")
    
    def export_docx(self, results, output_file, highlight=True, command_line='', username=None):
        """Export results to DOCX file with optional highlighting"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            # Create output directory if it doesn't exist
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = Document()
            
            if username:
                doc.add_heading(f'Comments by {username}', 0)
            else:
                doc.add_heading('YouTube Comment Search Results', 0)
            
            # Add command line at the top
            if command_line:
                p = doc.add_paragraph()
                p.add_run('Command: ').bold = True
                p.add_run(command_line)
                doc.add_paragraph()
            
            doc.add_paragraph(f'Total Results: {len(results)}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            for i, result in enumerate(results, 1):
                comment = result['comment']
                
                # Add result header
                doc.add_heading(f'Result {i}', level=1)
                
                # Add metadata
                p = doc.add_paragraph()
                p.add_run(f"Video ID: ").bold = True
                p.add_run(f"{result['video_id']}\n")
                p.add_run(f"Author: ").bold = True
                p.add_run(f"{comment.get('author', 'Unknown')} ({comment.get('author_id', 'N/A')})\n")
                p.add_run(f"Likes: ").bold = True
                p.add_run(f"{comment.get('like_count', 0)} | ")
                p.add_run(f"Verified: {('Yes' if comment.get('author_is_verified') else 'No')} | ")
                p.add_run(f"Uploader: {('Yes' if comment.get('author_is_uploader') else 'No')}\n")
                p.add_run(f"Posted: ").bold = True
                p.add_run(f"{comment.get('_time_text', 'Unknown')}\n")
                
                if result['relevance_score'] > 0:
                    p.add_run(f"Relevance Score: ").bold = True
                    p.add_run(f"{result['relevance_score']}\n")
                
                p.add_run(f"URL: ").bold = True
                p.add_run(f"https://www.youtube.com/watch?v={result['video_id']}\n")
                
                # Add comment text with highlighting
                comment_para = doc.add_paragraph()
                comment_para.add_run("Comment: ").bold = True
                
                text = comment.get('text', '')
                
                if highlight and result.get('matched_text'):
                    # Highlight matched portions in red
                    matched_terms = result['matched_text']
                    
                    # Create a regex pattern for all matched terms
                    patterns = [re.escape(term) for term in matched_terms]
                    combined_pattern = '|'.join(patterns)
                    
                    try:
                        pattern = re.compile(f'({combined_pattern})', re.IGNORECASE)
                        parts = pattern.split(text)
                        
                        for part in parts:
                            if part and any(part.lower() == term.lower() for term in matched_terms):
                                # This is a matched term - color it red
                                run = comment_para.add_run(part)
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            else:
                                # Regular text
                                comment_para.add_run(part)
                    except:
                        # Fallback: just add plain text
                        comment_para.add_run(text)
                else:
                    comment_para.add_run(text)
                
                doc.add_paragraph()  # Add spacing between results
            
            doc.save(output_file)
            print(f"✓ Results exported to DOCX: {output_file}")
            
        except ImportError:
            print(f"✗ Error: python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            print(f"✗ Error exporting DOCX results: {e}")
    
    def generate_author_stats(self, results):
        """Generate statistics about comment authors"""
        author_stats = defaultdict(lambda: {'count': 0, 'total_likes': 0, 'comments': []})
        
        for result in results:
            comment = result['comment']
            author = comment.get('author', 'Unknown')
            
            author_stats[author]['count'] += 1
            author_stats[author]['total_likes'] += comment.get('like_count', 0)
            author_stats[author]['comments'].append(comment.get('text', ''))
        
        print(f"\n{'='*60}")
        print("Author Statistics")
        print(f"{'='*60}\n")
        
        # Sort by comment count
        sorted_authors = sorted(author_stats.items(), 
                               key=lambda x: x[1]['count'], 
                               reverse=True)
        
        print(f"Total unique authors: {len(sorted_authors)}\n")
        print("Top 10 most active authors:")
        for i, (author, stats) in enumerate(sorted_authors[:10], 1):
            print(f"{i}. {author}")
            print(f"   Comments: {stats['count']} | Total likes: {stats['total_likes']}")
        print()

def main():
    from comment_search_help import get_usage_examples, get_arg_parser
    
    # Check if no arguments provided
    if len(sys.argv) == 1:
        print("\n" + "="*60)
        print(APP_NAME.center(60))
        print("="*60)
        print("\nError: No arguments provided\n")
        print("Usage: python comment_search.py -d DIRECTORY [OPTIONS]\n")
        print(get_usage_examples())
        print("\nFor more help, use: python comment_search.py -h\n")
        sys.exit(1)
    

    parser = get_arg_parser()
    parser.add_argument(
        "--no-color",
        action="store_true",
        help="Disable ANSI colors (for GUI output)"
    )
    args = parser.parse_args()



    USE_COLOR = not args.no_color

    # Create searcher instance
    searcher = CommentSearcher(args.dir)
    
    # Handle --most-active flag
    if args.most_active:
        top_users = searcher.show_most_active_users(args.most_active)
        
        if not top_users:
            return
        
        # If keywords provided, search within most active user(s) comments
        if args.keywords:
            all_results = []
            
            for author, stats in top_users:
                print(f"\nSearching in comments by: {author}")
                print(f"{'='*60}\n")
                
                # Extract comments from this user with keyword filter
                user_results = searcher.extract_user_comments(
                    username=author,
                    search_keywords=args.keywords,
                    search_mode=args.mode,
                    case_sensitive=args.case_sensitive,
                    min_likes=args.min_likes
                )
                
                all_results.extend(user_results)
            
            # Apply max results limit if specified
            if args.max_results and len(all_results) > args.max_results:
                all_results = all_results[:args.max_results]
            
            # Display results
            if all_results:
                searcher.display_results(
                    all_results,
                    highlight=not args.no_highlight,
                    keywords=args.keywords,
                    search_mode=args.mode,
                    case_sensitive=args.case_sensitive
                )
            
            # Auto-save results unless disabled
            if not args.no_save and all_results:
                output_dir = Path('search_results')
                output_dir.mkdir(exist_ok=True)
                
                dir_name = Path(args.dir).name
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Create descriptive filename
                if args.most_active == 1:
                    safe_username = "".join(c for c in top_users[0][0] if c.isalnum() or c in (' ', '-', '_')).strip()
                    safe_username = safe_username.replace(' ', '_')
                    base_filename = output_dir / f"{dir_name}_most_active_{safe_username}_{timestamp}"
                else:
                    base_filename = output_dir / f"{dir_name}_most_active_top{args.most_active}_{timestamp}"
                
                command_line = ' '.join(sys.argv)
                
                # Save in requested formats
                for export_format in args.export:
                    if export_format == 'json':
                        json_file = f"{base_filename}.json"
                        searcher.export_json(all_results, json_file, command_line)
                    elif export_format == 'docx':
                        docx_file = f"{base_filename}.docx"
                        searcher.export_docx(
                            all_results,
                            docx_file,
                            highlight=not args.no_docx_highlight,
                            command_line=command_line
                        )
                
                print()
        else:
            # No keywords - export user statistics
            if not args.no_save:
                output_dir = Path('search_results')
                output_dir.mkdir(exist_ok=True)
                
                dir_name = Path(args.dir).name
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_filename = output_dir / f"{dir_name}_most_active_top{args.most_active}_{timestamp}"
                
                command_line = ' '.join(sys.argv)
                
                # Save in requested formats
                for export_format in args.export:
                    if export_format == 'json':
                        json_file = f"{base_filename}.json"
                        searcher.export_active_users_json(top_users, json_file, command_line)
                    elif export_format == 'docx':
                        docx_file = f"{base_filename}.docx"
                        searcher.export_active_users_docx(top_users, docx_file, command_line)
                
                print()
        
        return
    
    # Handle --user flag
    if args.user:
        # Extract user comments, optionally with search keywords
        results = searcher.extract_user_comments(
            username=args.user,
            search_keywords=args.keywords if args.keywords else None,
            search_mode=args.mode,
            case_sensitive=args.case_sensitive,
            min_likes=args.min_likes
        )
        
        # Apply max results limit if specified
        if args.max_results and len(results) > args.max_results:
            results = results[:args.max_results]
        
        # Display results
        if results:
            searcher.display_results(
                results,
                highlight=not args.no_highlight,
                keywords=args.keywords if args.keywords else None,
                search_mode=args.mode,
                case_sensitive=args.case_sensitive
            )
        
        # Auto-save results unless disabled
        if not args.no_save and results:
            # Create output directory
            output_dir = Path('search_results')
            output_dir.mkdir(exist_ok=True)
            
            # Generate base filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_username = "".join(c for c in args.user if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_username = safe_username.replace(' ', '_')
            base_filename = output_dir / f"user_{safe_username}_{timestamp}"
            
            # Reconstruct command line
            command_line = ' '.join(sys.argv)
            
            # Save in requested formats
            for export_format in args.export:
                if export_format == 'json':
                    json_file = f"{base_filename}.json"
                    searcher.export_json(results, json_file, command_line, username=args.user)
                elif export_format == 'docx':
                    docx_file = f"{base_filename}.docx"
                    searcher.export_docx(
                        results,
                        docx_file,
                        highlight=not args.no_docx_highlight,
                        command_line=command_line,
                        username=args.user
                    )
            
            print()
        
        # Show author stats if requested (though for single user it's redundant)
        if args.stats and results:
            searcher.generate_author_stats(results)
        
        return
    
    # Regular keyword search (original functionality)
    if not args.keywords:
        print("\n✗ Error: Keywords are required for regular search")
        print("Use --most-active to find active users, or --user to extract user comments\n")
        sys.exit(1)
    
    # Perform search
    results = searcher.search_comments(
        keywords=args.keywords,
        search_mode=args.mode,
        case_sensitive=args.case_sensitive,
        min_likes=args.min_likes,
        max_results=args.max_results,
        sort_by=args.sort
    )
    
    # Display results
    searcher.display_results(
        results, 
        highlight=not args.no_highlight,
        keywords=args.keywords,
        search_mode=args.mode,
        case_sensitive=args.case_sensitive
    )
    
    # Auto-save results unless disabled
    if not args.no_save and results:
        # Create output directory
        output_dir = Path('search_results')
        output_dir.mkdir(exist_ok=True)
        
        # Generate base filename from directory name with timestamp
        dir_name = Path(args.dir).name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = output_dir / f"{dir_name}_{timestamp}"
        
        # Reconstruct command line
        command_line = ' '.join(sys.argv)
        
        # Save in requested formats
        for export_format in args.export:
            if export_format == 'json':
                json_file = f"{base_filename}.json"
                searcher.export_json(results, json_file, command_line)
            elif export_format == 'docx':
                docx_file = f"{base_filename}.docx"
                searcher.export_docx(
                    results, 
                    docx_file, 
                    highlight=not args.no_docx_highlight,
                    command_line=command_line
                )
        
        print()
    
    # Show author stats if requested
    if args.stats and results:
        searcher.generate_author_stats(results)

if __name__ == "__main__":
    main()