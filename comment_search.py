#!/usr/bin/env python3
"""
Search through downloaded YouTube comments for keywords and patterns
"""

import json
import argparse
import re
import sys
from pathlib import Path
from collections import defaultdict
from datetime import datetime

APP_NAME = "YouTube Comment Search"

# ANSI color codes for terminal highlighting
RED = '\033[91m'
RESET = '\033[0m'

class CommentSearcher:
    def __init__(self, comments_dir):
        self.comments_dir = Path(comments_dir)
        self.results = []
        
    def load_comments_from_file(self, filepath):
        """Load comments from a JSON file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                comments = json.load(f)
            return comments, filepath.name
        except Exception as e:
            print(f"✗ Error loading {filepath}: {e}")
            return [], None
    
    def search_simple(self, text, keywords, case_sensitive=False):
        """Simple keyword search (any keyword matches)"""
        if not case_sensitive:
            text = text.lower()
            keywords = [k.lower() for k in keywords]
        
        return any(keyword in text for keyword in keywords)
    
    def search_all_keywords(self, text, keywords, case_sensitive=False):
        """Search requiring ALL keywords to be present"""
        if not case_sensitive:
            text = text.lower()
            keywords = [k.lower() for k in keywords]
        
        return all(keyword in text for keyword in keywords)
    
    def search_phrase(self, text, phrase, case_sensitive=False):
        """Search for exact phrase"""
        if not case_sensitive:
            text = text.lower()
            phrase = phrase.lower()
        
        return phrase in text
    
    def search_regex(self, text, pattern):
        """Search using regular expression"""
        try:
            return bool(re.search(pattern, text, re.IGNORECASE))
        except re.error as e:
            print(f"✗ Invalid regex pattern: {e}")
            return False
    
    def highlight_matches(self, text, keywords, search_mode, case_sensitive=False, use_color=True):
        """Highlight matching keywords in text"""
        if not use_color:
            return text
        
        highlighted_text = text
        
        if search_mode == 'regex':
            # For regex, highlight actual matched portions
            try:
                pattern = re.compile(keywords[0], re.IGNORECASE if not case_sensitive else 0)
                matches = list(pattern.finditer(text))
                # Process matches in reverse to maintain string positions
                for match in reversed(matches):
                    start, end = match.span()
                    highlighted_text = (highlighted_text[:start] + 
                                      RED + highlighted_text[start:end] + RESET + 
                                      highlighted_text[end:])
            except re.error:
                pass
        elif search_mode == 'phrase':
            # Highlight the entire phrase
            phrase = ' '.join(keywords)
            if not case_sensitive:
                # Case-insensitive phrase highlighting
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                highlighted_text = pattern.sub(lambda m: RED + m.group(0) + RESET, highlighted_text)
            else:
                highlighted_text = highlighted_text.replace(phrase, RED + phrase + RESET)
        else:
            # Highlight individual keywords (any or all mode)
            for keyword in keywords:
                if not case_sensitive:
                    # Case-insensitive keyword highlighting
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    highlighted_text = pattern.sub(lambda m: RED + m.group(0) + RESET, highlighted_text)
                else:
                    highlighted_text = highlighted_text.replace(keyword, RED + keyword + RESET)
        
        return highlighted_text
    
    def get_plain_matches(self, text, keywords, search_mode, case_sensitive=False):
        """Get list of actual matched text portions (for DOCX highlighting)"""
        matches = []
        
        if search_mode == 'regex':
            try:
                pattern = re.compile(keywords[0], re.IGNORECASE if not case_sensitive else 0)
                for match in pattern.finditer(text):
                    matches.append(match.group(0))
            except re.error:
                pass
        elif search_mode == 'phrase':
            phrase = ' '.join(keywords)
            if not case_sensitive:
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                matches = [m.group(0) for m in pattern.finditer(text)]
            else:
                # Find all occurrences of exact phrase
                start = 0
                while True:
                    pos = text.find(phrase, start)
                    if pos == -1:
                        break
                    matches.append(text[pos:pos+len(phrase)])
                    start = pos + 1
        else:
            for keyword in keywords:
                if not case_sensitive:
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    matches.extend([m.group(0) for m in pattern.finditer(text)])
                else:
                    start = 0
                    while True:
                        pos = text.find(keyword, start)
                        if pos == -1:
                            break
                        matches.append(text[pos:pos+len(keyword)])
                        start = pos + 1
        
        return list(set(matches))  # Remove duplicates
    
    def calculate_relevance_score(self, comment, keywords, case_sensitive=False):
        """Calculate relevance score based on keyword frequency and position"""
        text = comment['text']
        if not case_sensitive:
            text_lower = text.lower()
            keywords_lower = [k.lower() for k in keywords]
        else:
            text_lower = text
            keywords_lower = keywords
        
        score = 0
        
        # Count keyword occurrences
        for keyword in keywords_lower:
            count = text_lower.count(keyword)
            score += count * 10
        
        # Bonus for keyword in first 50 characters
        if any(kw in text_lower[:50] for kw in keywords_lower):
            score += 5
        
        # Bonus for higher like count
        score += min(comment.get('like_count', 0), 50)
        
        # Bonus for uploader/verified authors
        if comment.get('author_is_uploader'):
            score += 20
        if comment.get('author_is_verified'):
            score += 10
        
        return score
    
    def search_comments(self, keywords, search_mode='any', case_sensitive=False, 
                       min_likes=0, max_results=None, sort_by='relevance'):
        """
        Search through all comment files
        
        search_mode: 'any', 'all', 'phrase', 'regex'
        sort_by: 'relevance', 'likes', 'date'
        """
        print(f"\n{'='*60}")
        print(f"Searching in: {self.comments_dir}")
        print(f"Keywords: {keywords}")
        print(f"Mode: {search_mode}, Case-sensitive: {case_sensitive}")
        print(f"Min likes: {min_likes}")
        print(f"{'='*60}\n")
        
        if not self.comments_dir.exists():
            print(f"✗ Error: Directory not found: {self.comments_dir}")
            return []
        
        json_files = list(self.comments_dir.glob("*.json"))
        
        if not json_files:
            print(f"✗ No JSON files found in {self.comments_dir}")
            return []
        
        print(f"Scanning {len(json_files)} comment file(s)...\n")
        
        results = []
        total_comments = 0
        
        for json_file in json_files:
            comments, filename = self.load_comments_from_file(json_file)
            total_comments += len(comments)
            
            for comment in comments:
                text = comment.get('text', '')
                if not text:
                    continue
                
                # Apply like filter
                if comment.get('like_count', 0) < min_likes:
                    continue
                
                # Apply search based on mode
                match = False
                if search_mode == 'any':
                    match = self.search_simple(text, keywords, case_sensitive)
                elif search_mode == 'all':
                    match = self.search_all_keywords(text, keywords, case_sensitive)
                elif search_mode == 'phrase':
                    match = self.search_phrase(text, ' '.join(keywords), case_sensitive)
                elif search_mode == 'regex':
                    match = self.search_regex(text, keywords[0])
                
                if match:
                    result = {
                        'comment': comment,
                        'source_file': filename,
                        'video_id': filename.split('_')[0] if '_' in filename else 'unknown',
                        'relevance_score': self.calculate_relevance_score(comment, keywords, case_sensitive),
                        'matched_text': self.get_plain_matches(text, keywords, search_mode, case_sensitive)
                    }
                    results.append(result)
        
        print(f"✓ Scanned {total_comments} total comments")
        print(f"✓ Found {len(results)} matching comments\n")
        
        # Sort results
        if sort_by == 'relevance':
            results.sort(key=lambda x: x['relevance_score'], reverse=True)
        elif sort_by == 'likes':
            results.sort(key=lambda x: x['comment'].get('like_count', 0), reverse=True)
        elif sort_by == 'date':
            results.sort(key=lambda x: x['comment'].get('timestamp', 0), reverse=True)
        
        # Limit results if specified
        if max_results:
            results = results[:max_results]
        
        return results
    
    def display_results(self, results, highlight=True, keywords=None, search_mode='any', 
                       case_sensitive=False):
        """Display search results"""
        if not results:
            print("No results found.")
            return
        
        print(f"{'='*60}")
        print(f"Top {len(results)} Results")
        print(f"{'='*60}\n")
        
        for i, result in enumerate(results, 1):
            comment = result['comment']
            
            print(f"[{i}] Video ID: {result['video_id']}")
            print(f"    Author: {comment.get('author', 'Unknown')} ({comment.get('author_id', 'N/A')})")
            print(f"    Likes: {comment.get('like_count', 0)} | "
                  f"Verified: {'Yes' if comment.get('author_is_verified') else 'No'} | "
                  f"Uploader: {'Yes' if comment.get('author_is_uploader') else 'No'}")
            print(f"    Posted: {comment.get('_time_text', 'Unknown')}")
            print(f"    Relevance Score: {result['relevance_score']}")
            
            # Display comment text (never truncated)
            text = comment.get('text', '')
            if highlight and keywords:
                text = self.highlight_matches(text, keywords, search_mode, case_sensitive, use_color=True)
            print(f"    Comment: {text}")
            
            print(f"    URL: https://www.youtube.com/watch?v={result['video_id']}")
            print()
    
    def export_json(self, results, output_file, command_line):
        """Export results to JSON file"""
        try:
            # Create output directory if it doesn't exist
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            export_data = {
                'command': command_line,
                'generated': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'total_results': len(results),
                'results': []
            }
            
            for result in results:
                export_data['results'].append({
                    'video_id': result['video_id'],
                    'source_file': result['source_file'],
                    'relevance_score': result['relevance_score'],
                    'matched_text': result.get('matched_text', []),
                    'comment': result['comment']
                })
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            print(f"✓ Results exported to JSON: {output_file}")
        except Exception as e:
            print(f"✗ Error exporting JSON results: {e}")
    
    def export_docx(self, results, output_file, highlight=True, command_line=''):
        """Export results to DOCX file with optional highlighting"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            # Create output directory if it doesn't exist
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = Document()
            doc.add_heading('YouTube Comment Search Results', 0)
            
            # Add command line at the top
            if command_line:
                p = doc.add_paragraph()
                p.add_run('Command: ').bold = True
                p.add_run(command_line)
                doc.add_paragraph()
            
            doc.add_paragraph(f'Total Results: {len(results)}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            for i, result in enumerate(results, 1):
                comment = result['comment']
                
                # Add result header
                doc.add_heading(f'Result {i}', level=1)
                
                # Add metadata
                p = doc.add_paragraph()
                p.add_run(f"Video ID: ").bold = True
                p.add_run(f"{result['video_id']}\n")
                p.add_run(f"Author: ").bold = True
                p.add_run(f"{comment.get('author', 'Unknown')} ({comment.get('author_id', 'N/A')})\n")
                p.add_run(f"Likes: ").bold = True
                p.add_run(f"{comment.get('like_count', 0)} | ")
                p.add_run(f"Verified: {('Yes' if comment.get('author_is_verified') else 'No')} | ")
                p.add_run(f"Uploader: {('Yes' if comment.get('author_is_uploader') else 'No')}\n")
                p.add_run(f"Posted: ").bold = True
                p.add_run(f"{comment.get('_time_text', 'Unknown')}\n")
                p.add_run(f"Relevance Score: ").bold = True
                p.add_run(f"{result['relevance_score']}\n")
                p.add_run(f"URL: ").bold = True
                p.add_run(f"https://www.youtube.com/watch?v={result['video_id']}\n")
                
                # Add comment text with highlighting
                comment_para = doc.add_paragraph()
                comment_para.add_run("Comment: ").bold = True
                
                text = comment.get('text', '')
                
                if highlight and result.get('matched_text'):
                    # Highlight matched portions in red
                    matched_terms = result['matched_text']
                    
                    # Create a regex pattern for all matched terms
                    patterns = [re.escape(term) for term in matched_terms]
                    combined_pattern = '|'.join(patterns)
                    
                    try:
                        pattern = re.compile(f'({combined_pattern})', re.IGNORECASE)
                        parts = pattern.split(text)
                        
                        for part in parts:
                            if part and any(part.lower() == term.lower() for term in matched_terms):
                                # This is a matched term - color it red
                                run = comment_para.add_run(part)
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            else:
                                # Regular text
                                comment_para.add_run(part)
                    except:
                        # Fallback: just add plain text
                        comment_para.add_run(text)
                else:
                    comment_para.add_run(text)
                
                doc.add_paragraph()  # Add spacing between results
            
            doc.save(output_file)
            print(f"✓ Results exported to DOCX: {output_file}")
            
        except ImportError:
            print(f"✗ Error: python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            print(f"✗ Error exporting DOCX results: {e}")
    
    def generate_author_stats(self, results):
        """Generate statistics about comment authors"""
        author_stats = defaultdict(lambda: {'count': 0, 'total_likes': 0, 'comments': []})
        
        for result in results:
            comment = result['comment']
            author = comment.get('author', 'Unknown')
            
            author_stats[author]['count'] += 1
            author_stats[author]['total_likes'] += comment.get('like_count', 0)
            author_stats[author]['comments'].append(comment.get('text', ''))
        
        print(f"\n{'='*60}")
        print("Author Statistics")
        print(f"{'='*60}\n")
        
        # Sort by comment count
        sorted_authors = sorted(author_stats.items(), 
                               key=lambda x: x[1]['count'], 
                               reverse=True)
        
        print(f"Total unique authors: {len(sorted_authors)}\n")
        print("Top 10 most active authors:")
        for i, (author, stats) in enumerate(sorted_authors[:10], 1):
            print(f"{i}. {author}")
            print(f"   Comments: {stats['count']} | Total likes: {stats['total_likes']}")
        print()

def main():
    # Check if no arguments provided
    if len(sys.argv) == 1:
        print("\n" + "="*60)
        print(APP_NAME.center(60))
        print("="*60)
        print("\nError: No keywords provided\n")
        print("Usage: python comment_search.py KEYWORDS [OPTIONS]\n")
        print("Examples:")
        print('  python comment_search.py "deep web" "murder" "photo"')
        print('  python comment_search.py "murder" -m all --min-likes 10')
        print('  python comment_search.py "deep web" --no-highlight')
        print('  python comment_search.py "murder" --export docx')
        print("\nFor more help, use: python comment_search.py -h\n")
        sys.exit(1)
    
    parser = argparse.ArgumentParser(
        description='Search through YouTube comment JSON files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Search for any keyword (saves JSON by default)
  %(prog)s "deep web" "murder" "photo"
  
  # Search requiring ALL keywords
  %(prog)s "deep web" "murder" -m all
  
  # Search for exact phrase
  %(prog)s "deep web murder" -m phrase
  
  # Disable highlighting in terminal
  %(prog)s "murder" --no-highlight
  
  # Export to DOCX with highlighting
  %(prog)s "murder" --export docx
  
  # Export to both JSON and DOCX
  %(prog)s "murder" --export json docx
  
  # Export to DOCX without highlighting
  %(prog)s "murder" --export docx --no-docx-highlight
  
  # Only show comments with 10+ likes
  %(prog)s "murder" --min-likes 10
  
  # Disable auto-save
  %(prog)s "murder" --no-save
  
  # Show author statistics
  %(prog)s "murder" --stats
  
  # Search with regex pattern
  %(prog)s "murder(ed|ing|s)?" -m regex
  
  # Custom comments directory
  %(prog)s "murder" -d ./my_comments
        """
    )
    
    parser.add_argument(
        'keywords',
        nargs='+',
        help='Keywords or phrases to search for'
    )
    
    parser.add_argument(
        '-d', '--dir',
        default='comment_sections',
        help='Directory containing comment JSON files (default: comment_sections)'
    )
    
    parser.add_argument(
        '-m', '--mode',
        choices=['any', 'all', 'phrase', 'regex'],
        default='any',
        help='Search mode: any=match any keyword, all=match all keywords, phrase=exact phrase, regex=regular expression (default: any)'
    )
    
    parser.add_argument(
        '--case-sensitive',
        action='store_true',
        help='Enable case-sensitive search'
    )
    
    parser.add_argument(
        '--min-likes',
        type=int,
        default=0,
        help='Minimum number of likes (default: 0)'
    )
    
    parser.add_argument(
        '-n', '--max-results',
        type=int,
        default=None,
        help='Maximum number of results to display (default: all)'
    )
    
    parser.add_argument(
        '-s', '--sort',
        choices=['relevance', 'likes', 'date'],
        default='relevance',
        help='Sort results by: relevance, likes, or date (default: relevance)'
    )
    
    parser.add_argument(
        '--no-highlight',
        action='store_true',
        help='Disable keyword highlighting in terminal output'
    )
    
    parser.add_argument(
        '--no-save',
        action='store_true',
        help='Disable automatic saving of results'
    )
    
    parser.add_argument(
        '--export',
        nargs='+',
        choices=['json', 'docx'],
        default=['json'],
        help='Export formats: json, docx (default: json). Can specify multiple formats.'
    )
    
    parser.add_argument(
        '--no-docx-highlight',
        action='store_true',
        help='Disable keyword highlighting in DOCX export'
    )
    
    parser.add_argument(
        '--stats',
        action='store_true',
        help='Show author statistics'
    )
    
    args = parser.parse_args()
    
    # Create searcher instance
    searcher = CommentSearcher(args.dir)
    
    # Perform search
    results = searcher.search_comments(
        keywords=args.keywords,
        search_mode=args.mode,
        case_sensitive=args.case_sensitive,
        min_likes=args.min_likes,
        max_results=args.max_results,
        sort_by=args.sort
    )
    
    # Display results
    searcher.display_results(
        results, 
        highlight=not args.no_highlight,
        keywords=args.keywords,
        search_mode=args.mode,
        case_sensitive=args.case_sensitive
    )
    
    # Auto-save results unless disabled
    if not args.no_save and results:
        # Create output directory
        output_dir = Path('search_results')
        output_dir.mkdir(exist_ok=True)
        
        # Generate base filename from directory name with timestamp
        dir_name = Path(args.dir).name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = output_dir / f"{dir_name}_{timestamp}"
        
        # Reconstruct command line
        command_line = ' '.join(sys.argv)
        
        # Save in requested formats
        for export_format in args.export:
            if export_format == 'json':
                json_file = f"{base_filename}.json"
                searcher.export_json(results, json_file, command_line)
            elif export_format == 'docx':
                docx_file = f"{base_filename}.docx"
                searcher.export_docx(
                    results, 
                    docx_file, 
                    highlight=not args.no_docx_highlight,
                    command_line=command_line
                )
        
        print()
    
    # Show author stats if requested
    if args.stats and results:
        searcher.generate_author_stats(results)

if __name__ == "__main__":
    main()